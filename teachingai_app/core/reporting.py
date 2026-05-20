from __future__ import annotations

import json
import ast

from html import escape
from io import BytesIO
from datetime import datetime

from .models import SimulationReport


def format_revised_lesson_plan(lesson_plan: str) -> str:
    """将修订后教案从JSON/Python字典格式转换为友好的文本格式"""
    if not lesson_plan.strip():
        return "（模型未返回修订后教案内容）"
    
    # 先尝试标准JSON解析
    try:
        parsed = json.loads(lesson_plan)
        return _format_json_plan(parsed)
    except json.JSONDecodeError:
        pass
    
    # 再尝试Python字典格式解析
    try:
        parsed = ast.literal_eval(lesson_plan)
        return _format_json_plan(parsed)
    except (SyntaxError, ValueError):
        return lesson_plan


def _format_json_plan(data) -> str:
    """递归格式化JSON数据为文本格式"""
    if isinstance(data, dict):
        result = []
        for key, value in data.items():
            formatted_value = _format_json_plan(value)
            if isinstance(value, (list, dict)):
                result.append(f"## {key}")
                result.append(formatted_value)
            else:
                result.append(f"**{key}**：{formatted_value}")
        return "\n\n".join(result)
    elif isinstance(data, list):
        result = []
        for i, item in enumerate(data, 1):
            formatted_item = _format_json_plan(item)
            if isinstance(item, dict):
                result.append(f"{i}. {formatted_item}")
            else:
                result.append(f"{i}. {formatted_item}")
        return "\n".join(result)
    else:
        return str(data)


def _mode_label(report: SimulationReport) -> str:
    return "深度思考模式（多轮讨论与裁决预演）" if report.analysis_mode == "deep" else "快速模式"


def _append_deep_markdown(lines: list[str], report: SimulationReport) -> None:
    if report.analysis_mode != "deep":
        return

    lines.append("## 深度预演：课堂讲稿")
    lines.append(report.teacher_script.strip() or "（未生成讲稿）")
    lines.append("")

    if report.lesson_modules:
        lines.append("## 深度预演：分模块互动记录")
        for module in sorted(report.lesson_modules, key=lambda m: m.order):
            lines.append(f"### 模块 {module.order}：{module.title}")
            lines.append(module.teacher_script.strip() or "")
            lines.append("")
            module_deliberation = next(
                (d for d in report.module_deliberations if d.module_id == module.module_id),
                None,
            )
            if module_deliberation is not None:
                lines.append("#### 讨论裁决结果（这一模块里学生讨论后，系统给出的最终归纳）")

                lines.append("- 模块共识（大家基本都同意的判断，可信度最高）:")
                if module_deliberation.consensus:
                    lines.extend([f"  - {idx}. {item}" for idx, item in enumerate(module_deliberation.consensus, start=1)])
                else:
                    lines.append("  - 1. 本模块暂无明确共识。")

                lines.append("- 即时调整（老师这节课当下就可以立刻改的做法）:")
                if module_deliberation.teaching_adjustments:
                    lines.extend(
                        [
                            f"  - {idx}. {item}"
                            for idx, item in enumerate(module_deliberation.teaching_adjustments, start=1)
                        ]
                    )
                else:
                    lines.append("  - 1. 本模块暂无即时调整建议。")

                lines.append("- 模块分歧（大家意见不一致的地方，说明这里还不确定）:")
                if module_deliberation.disagreements:
                    lines.extend(
                        [f"  - {idx}. {item}" for idx, item in enumerate(module_deliberation.disagreements, start=1)]
                    )
                else:
                    lines.append("  - 1. 本模块暂无显著分歧。")

                lines.append("- 记忆更新（学生状态变化，会带到下一模块继续影响表现）:")
                if module_deliberation.memory_updates:
                    lines.extend(
                        [f"  - {idx}. {item}" for idx, item in enumerate(module_deliberation.memory_updates, start=1)]
                    )
                else:
                    lines.append("  - 1. 本模块暂无记忆状态更新。")
            module_items = [
                i
                for i in report.module_interactions
                if i.module_id == module.module_id
            ]
            for inter in module_items:
                lines.append(
                    f"#### {inter.profile_name}（参与度 {inter.engagement}，置信度 {inter.confidence_score}/100）"
                )
                if inter.verbal_response:
                    lines.append(f"- 课堂发言: {inter.verbal_response}")
                if inter.consistency_note:
                    lines.append(f"- 一致性说明: {inter.consistency_note}")
                for point in inter.confusion_points:
                    lines.append(f"- 困惑: {point}")
                for q in inter.likely_questions:
                    lines.append(f"- 提问: {q}")
                for err in inter.error_predictions:
                    lines.append(f"- 错误: {err}")
            lines.append("")


def _append_deep_html(parts: list[str], report: SimulationReport) -> None:
    if report.analysis_mode != "deep":
        return

    parts.append("<section><h2>深度预演：课堂讲稿</h2>")
    script = report.teacher_script.strip() or "（未生成讲稿）"
    for paragraph in script.splitlines():
        text = paragraph.strip()
        if text:
            parts.append(f"<p>{escape(text)}</p>")
    parts.append("</section>")

    if report.lesson_modules:
        parts.append("<section><h2>深度预演：分模块互动记录</h2>")
        for module in sorted(report.lesson_modules, key=lambda m: m.order):
            parts.append(f"<h3>模块 {module.order}：{escape(module.title)}</h3>")
            for paragraph in (module.teacher_script or "").splitlines():
                text = paragraph.strip()
                if text:
                    parts.append(f"<p>{escape(text)}</p>")
            module_deliberation = next(
                (d for d in report.module_deliberations if d.module_id == module.module_id),
                None,
            )
            if module_deliberation is not None:
                parts.append("<h4>讨论裁决结果（这一模块里学生讨论后，系统给出的最终归纳）</h4>")

                parts.append("<p><strong>模块共识（大家基本都同意的判断，可信度最高）</strong></p><ul>")
                if module_deliberation.consensus:
                    for idx, item in enumerate(module_deliberation.consensus, start=1):
                        parts.append(f"<li>{idx}. {escape(item)}</li>")
                else:
                    parts.append("<li>1. 本模块暂无明确共识。</li>")
                parts.append("</ul>")

                parts.append("<p><strong>即时调整（老师这节课当下就可以立刻改的做法）</strong></p><ul>")
                if module_deliberation.teaching_adjustments:
                    for idx, item in enumerate(module_deliberation.teaching_adjustments, start=1):
                        parts.append(f"<li>{idx}. {escape(item)}</li>")
                else:
                    parts.append("<li>1. 本模块暂无即时调整建议。</li>")
                parts.append("</ul>")

                parts.append("<p><strong>模块分歧（大家意见不一致的地方，说明这里还不确定）</strong></p><ul>")
                if module_deliberation.disagreements:
                    for idx, item in enumerate(module_deliberation.disagreements, start=1):
                        parts.append(f"<li>{idx}. {escape(item)}</li>")
                else:
                    parts.append("<li>1. 本模块暂无显著分歧。</li>")
                parts.append("</ul>")

                parts.append("<p><strong>记忆更新（学生状态变化，会带到下一模块继续影响表现）</strong></p><ul>")
                if module_deliberation.memory_updates:
                    for idx, item in enumerate(module_deliberation.memory_updates, start=1):
                        parts.append(f"<li>{idx}. {escape(item)}</li>")
                else:
                    parts.append("<li>1. 本模块暂无记忆状态更新。</li>")
                parts.append("</ul>")
            module_items = [
                i for i in report.module_interactions if i.module_id == module.module_id
            ]
            for inter in module_items:
                parts.append(
                    f"<h4>{escape(inter.profile_name)}"
                    f"<span class='tag'>参与度 {escape(inter.engagement)}</span>"
                    f"<span class='tag'>置信度 {inter.confidence_score}/100</span></h4>"
                )
                if inter.verbal_response:
                    parts.append(f"<p><strong>课堂发言:</strong> {escape(inter.verbal_response)}</p>")
                if inter.consistency_note:
                    parts.append(f"<p><strong>一致性说明:</strong> {escape(inter.consistency_note)}</p>")
                parts.append("<ul>")
                for point in inter.confusion_points:
                    parts.append(f"<li>困惑: {escape(point)}</li>")
                for q in inter.likely_questions:
                    parts.append(f"<li>提问: {escape(q)}</li>")
                for err in inter.error_predictions:
                    parts.append(f"<li>错误: {escape(err)}</li>")
                parts.append("</ul>")
        parts.append("</section>")


def _append_deep_docx(doc, report: SimulationReport) -> None:
    if report.analysis_mode != "deep":
        return

    doc.add_heading("深度预演：课堂讲稿", level=2)
    for paragraph in (report.teacher_script or "（未生成讲稿）").splitlines():
        text = paragraph.strip()
        if text:
            doc.add_paragraph(text)

    if report.lesson_modules:
        doc.add_heading("深度预演：分模块互动记录", level=2)
        for module in sorted(report.lesson_modules, key=lambda m: m.order):
            doc.add_heading(f"模块 {module.order}：{module.title}", level=3)
            for paragraph in (module.teacher_script or "").splitlines():
                text = paragraph.strip()
                if text:
                    doc.add_paragraph(text)
            module_deliberation = next(
                (d for d in report.module_deliberations if d.module_id == module.module_id),
                None,
            )
            if module_deliberation is not None:
                doc.add_heading("讨论裁决结果（这一模块里学生讨论后，系统给出的最终归纳）", level=4)

                doc.add_paragraph("模块共识（大家基本都同意的判断，可信度最高）")
                if module_deliberation.consensus:
                    for idx, item in enumerate(module_deliberation.consensus, start=1):
                        doc.add_paragraph(f"{idx}. {item}", style="List Bullet")
                else:
                    doc.add_paragraph("1. 本模块暂无明确共识。", style="List Bullet")

                doc.add_paragraph("即时调整（老师这节课当下就可以立刻改的做法）")
                if module_deliberation.teaching_adjustments:
                    for idx, item in enumerate(module_deliberation.teaching_adjustments, start=1):
                        doc.add_paragraph(f"{idx}. {item}", style="List Bullet")
                else:
                    doc.add_paragraph("1. 本模块暂无即时调整建议。", style="List Bullet")

                doc.add_paragraph("模块分歧（大家意见不一致的地方，说明这里还不确定）")
                if module_deliberation.disagreements:
                    for idx, item in enumerate(module_deliberation.disagreements, start=1):
                        doc.add_paragraph(f"{idx}. {item}", style="List Bullet")
                else:
                    doc.add_paragraph("1. 本模块暂无显著分歧。", style="List Bullet")

                doc.add_paragraph("记忆更新（学生状态变化，会带到下一模块继续影响表现）")
                if module_deliberation.memory_updates:
                    for idx, item in enumerate(module_deliberation.memory_updates, start=1):
                        doc.add_paragraph(f"{idx}. {item}", style="List Bullet")
                else:
                    doc.add_paragraph("1. 本模块暂无记忆状态更新。", style="List Bullet")
            module_items = [
                i for i in report.module_interactions if i.module_id == module.module_id
            ]
            for inter in module_items:
                doc.add_heading(
                    f"{inter.profile_name}（参与度 {inter.engagement}，置信度 {inter.confidence_score}/100）",
                    level=4,
                )
                if inter.verbal_response:
                    doc.add_paragraph(f"课堂发言: {inter.verbal_response}")
                if inter.consistency_note:
                    doc.add_paragraph(f"一致性说明: {inter.consistency_note}")
                for point in inter.confusion_points:
                    doc.add_paragraph(f"困惑: {point}", style="List Bullet")
                for q in inter.likely_questions:
                    doc.add_paragraph(f"提问: {q}", style="List Bullet")
                for err in inter.error_predictions:
                    doc.add_paragraph(f"错误: {err}", style="List Bullet")


def to_markdown(report: SimulationReport) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines: list[str] = []

    lines.append("# AI虚拟学生教学预演报告")
    lines.append("")
    lines.append(f"- 生成时间: {now}")
    lines.append(f"- 学科: {report.subject}")
    lines.append(f"- 课题: {report.lesson_topic}")
    lines.append(f"- 年级: {report.grade}")
    lines.append(f"- 分析模式: {_mode_label(report)}")
    lines.append("")

    _append_deep_markdown(lines, report)

    lines.append("## 关键知识点")
    for point in report.extracted_key_points:
        lines.append(f"- {point}")
    lines.append("")

    lines.append("## 学生反应模拟")
    for reaction in report.reactions:
        lines.append(f"### {reaction.profile_name}")
        lines.append(f"- 参与度预估: {reaction.engagement}")
        lines.append("- 可能困惑点:")
        for point in reaction.confusion_points:
            lines.append(f"  - {point}")
        lines.append("- 可能提问:")
        for q in reaction.likely_questions:
            lines.append(f"  - {q}")
        lines.append("- 可能错误:")
        for e in reaction.error_predictions:
            lines.append(f"  - {e}")

    if report.difficulty:
        lines.append("")
        lines.append("## 难度评估")
        lines.append(f"- 综合难度: {report.difficulty.overall_level}")
        lines.append(
            f"- 认知负荷: {report.difficulty.cognitive_load_score}/10 | "
            f"步骤复杂度: {report.difficulty.step_complexity_score}/10 | "
            f"概念跨度: {report.difficulty.concept_span_score}/10"
        )
        for item in report.difficulty.rationale:
            lines.append(f"- {item}")

    if report.confidence:
        lines.append("")
        lines.append("## 结果置信度")
        lines.append(f"- 置信等级: {report.confidence.overall_level}")
        lines.append(f"- 置信分数: {report.confidence.overall_score}/100")
        for item in report.confidence.rationale:
            lines.append(f"- {item}")
        for item in report.confidence.profile_confidence:
            lines.append(f"- {item}")

    lines.append("")
    lines.append("## 教学优化建议")
    for idx, s in enumerate(report.suggestions, start=1):
        lines.append(f"### 建议 {idx}（优先级: {s.priority}）")
        lines.append(f"- 问题: {s.issue}")
        lines.append(f"- 建议: {s.suggestion}")
        lines.append(f"- 预期效果: {s.expected_impact}")

    return "\n".join(lines)


def to_html(report: SimulationReport) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    parts: list[str] = []
    parts.append("<!DOCTYPE html>")
    parts.append("<html lang='zh-CN'>")
    parts.append("<head>")
    parts.append("<meta charset='utf-8'>")
    parts.append("<meta name='viewport' content='width=device-width, initial-scale=1'>")
    parts.append("<title>AI虚拟学生教学预演报告</title>")
    parts.append(
        "<style>body{font-family:'Microsoft YaHei',sans-serif;max-width:960px;margin:32px auto;padding:0 16px;line-height:1.7;color:#1f2937;}"
        "h1,h2,h3{color:#111827;}"
        "section{margin-bottom:28px;padding:20px;border:1px solid #e5e7eb;border-radius:14px;background:#fcfcfd;}"
        "ul{padding-left:20px;}"
        ".meta li{list-style:none;margin:4px 0;}"
        ".tag{display:inline-block;padding:2px 10px;border-radius:999px;background:#eef2ff;color:#3730a3;font-size:12px;margin-left:8px;}"
        "</style>"
    )
    parts.append("</head>")
    parts.append("<body>")
    parts.append("<h1>AI虚拟学生教学预演报告</h1>")
    parts.append("<section>")
    parts.append("<ul class='meta'>")
    parts.append(f"<li><strong>生成时间:</strong> {escape(now)}</li>")
    parts.append(f"<li><strong>学科:</strong> {escape(report.subject)}</li>")
    parts.append(f"<li><strong>课题:</strong> {escape(report.lesson_topic)}</li>")
    parts.append(f"<li><strong>年级:</strong> {escape(report.grade)}</li>")
    parts.append(f"<li><strong>分析模式:</strong> {escape(_mode_label(report))}</li>")
    parts.append("</ul>")
    parts.append("</section>")

    _append_deep_html(parts, report)

    parts.append("<section><h2>关键知识点</h2><ul>")
    for point in report.extracted_key_points:
        parts.append(f"<li>{escape(point)}</li>")
    parts.append("</ul></section>")

    parts.append("<section><h2>学生反应模拟</h2>")
    for reaction in report.reactions:
        parts.append(
            f"<h3>{escape(reaction.profile_name)}<span class='tag'>参与度 {escape(reaction.engagement)}</span></h3>"
        )
        parts.append("<p><strong>可能困惑点</strong></p><ul>")
        for point in reaction.confusion_points:
            parts.append(f"<li>{escape(point)}</li>")
        parts.append("</ul><p><strong>可能提问</strong></p><ul>")
        for question in reaction.likely_questions:
            parts.append(f"<li>{escape(question)}</li>")
        parts.append("</ul><p><strong>可能错误</strong></p><ul>")
        for item in reaction.error_predictions:
            parts.append(f"<li>{escape(item)}</li>")
        parts.append("</ul>")
    parts.append("</section>")

    if report.difficulty:
        parts.append("<section><h2>难度评估</h2>")
        parts.append(f"<p><strong>综合难度:</strong> {escape(report.difficulty.overall_level)}</p>")
        parts.append(
            f"<p><strong>认知负荷:</strong> {report.difficulty.cognitive_load_score}/10 | "
            f"<strong>步骤复杂度:</strong> {report.difficulty.step_complexity_score}/10 | "
            f"<strong>概念跨度:</strong> {report.difficulty.concept_span_score}/10</p>"
        )
        parts.append("<ul>")
        for item in report.difficulty.rationale:
            parts.append(f"<li>{escape(item)}</li>")
        parts.append("</ul></section>")

    if report.confidence:
        parts.append("<section><h2>结果置信度</h2>")
        parts.append(f"<p><strong>置信等级:</strong> {escape(report.confidence.overall_level)}</p>")
        parts.append(f"<p><strong>置信分数:</strong> {report.confidence.overall_score}/100</p>")
        parts.append("<ul>")
        for item in report.confidence.rationale:
            parts.append(f"<li>{escape(item)}</li>")
        for item in report.confidence.profile_confidence:
            parts.append(f"<li>{escape(item)}</li>")
        parts.append("</ul></section>")

    parts.append("<section><h2>教学优化建议</h2>")
    for idx, suggestion in enumerate(report.suggestions, start=1):
        parts.append(
            f"<h3>建议 {idx}<span class='tag'>优先级 {escape(suggestion.priority)}</span></h3>"
        )
        parts.append(f"<p><strong>问题:</strong> {escape(suggestion.issue)}</p>")
        parts.append(f"<p><strong>建议:</strong> {escape(suggestion.suggestion)}</p>")
        parts.append(f"<p><strong>预期效果:</strong> {escape(suggestion.expected_impact)}</p>")
    parts.append("</section>")

    parts.append("</body></html>")

    return "".join(parts)


def to_docx_bytes(report: SimulationReport) -> bytes:
    from docx import Document

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc = Document()
    doc.add_heading("AI虚拟学生教学预演报告", level=1)
    doc.add_paragraph(f"生成时间: {now}")
    doc.add_paragraph(f"学科: {report.subject}")
    doc.add_paragraph(f"课题: {report.lesson_topic}")
    doc.add_paragraph(f"年级: {report.grade}")
    doc.add_paragraph(f"分析模式: {_mode_label(report)}")

    _append_deep_docx(doc, report)

    doc.add_heading("关键知识点", level=2)
    for point in report.extracted_key_points:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("学生反应模拟", level=2)
    for reaction in report.reactions:
        doc.add_heading(reaction.profile_name, level=3)
        doc.add_paragraph(f"参与度预估: {reaction.engagement}")
        doc.add_paragraph("可能困惑点")
        for point in reaction.confusion_points:
            doc.add_paragraph(point, style="List Bullet")
        doc.add_paragraph("可能提问")
        for question in reaction.likely_questions:
            doc.add_paragraph(question, style="List Bullet")
        doc.add_paragraph("可能错误")
        for item in reaction.error_predictions:
            doc.add_paragraph(item, style="List Bullet")

    if report.difficulty:
        doc.add_heading("难度评估", level=2)
        doc.add_paragraph(f"综合难度: {report.difficulty.overall_level}")
        doc.add_paragraph(
            f"认知负荷: {report.difficulty.cognitive_load_score}/10 | "
            f"步骤复杂度: {report.difficulty.step_complexity_score}/10 | "
            f"概念跨度: {report.difficulty.concept_span_score}/10"
        )
        for item in report.difficulty.rationale:
            doc.add_paragraph(item, style="List Bullet")

    if report.confidence:
        doc.add_heading("结果置信度", level=2)
        doc.add_paragraph(f"置信等级: {report.confidence.overall_level}")
        doc.add_paragraph(f"置信分数: {report.confidence.overall_score}/100")
        for item in report.confidence.rationale:
            doc.add_paragraph(item, style="List Bullet")
        for item in report.confidence.profile_confidence:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("教学优化建议", level=2)
    for idx, suggestion in enumerate(report.suggestions, start=1):
        doc.add_heading(f"建议 {idx}（优先级: {suggestion.priority}）", level=3)
        doc.add_paragraph(f"问题: {suggestion.issue}")
        doc.add_paragraph(f"建议: {suggestion.suggestion}")
        doc.add_paragraph(f"预期效果: {suggestion.expected_impact}")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def to_markdown_revised_plan(report: SimulationReport) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines: list[str] = []
    lines.append("# AI修订后教案")
    lines.append("")
    lines.append(f"- 生成时间: {now}")
    lines.append(f"- 学科: {report.subject}")
    lines.append(f"- 课题: {report.lesson_topic}")
    lines.append(f"- 年级: {report.grade}")

    if report.lesson_plan_change_summary:
        lines.append("")
        lines.append("## 教案修改摘要")
        for item in report.lesson_plan_change_summary:
            lines.append(f"- {item}")

    if report.original_lesson_material.strip():
        lines.append("")
        lines.append("## 原始教案")
        lines.append(report.original_lesson_material)

    lines.append("")
    lines.append("## 修订后教案")
    lines.append(format_revised_lesson_plan(report.revised_lesson_plan))
    return "\n".join(lines)


def to_html_revised_plan(report: SimulationReport) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    parts: list[str] = []
    parts.append("<!DOCTYPE html><html lang='zh-CN'><head><meta charset='utf-8'>")
    parts.append("<meta name='viewport' content='width=device-width, initial-scale=1'>")
    parts.append("<title>AI修订后教案</title>")
    parts.append(
        "<style>body{font-family:'Microsoft YaHei',sans-serif;max-width:960px;margin:32px auto;padding:0 16px;line-height:1.7;color:#1f2937;}"
        "h1,h2{color:#111827;}section{margin-bottom:24px;padding:18px;border:1px solid #e5e7eb;border-radius:12px;background:#fcfcfd;}ul{padding-left:20px;}"
        "</style></head><body>"
    )
    parts.append("<h1>AI修订后教案</h1><section><ul>")
    parts.append(f"<li><strong>生成时间:</strong> {escape(now)}</li>")
    parts.append(f"<li><strong>学科:</strong> {escape(report.subject)}</li>")
    parts.append(f"<li><strong>课题:</strong> {escape(report.lesson_topic)}</li>")
    parts.append(f"<li><strong>年级:</strong> {escape(report.grade)}</li></ul></section>")

    if report.lesson_plan_change_summary:
        parts.append("<section><h2>教案修改摘要</h2><ul>")
        for item in report.lesson_plan_change_summary:
            parts.append(f"<li>{escape(item)}</li>")
        parts.append("</ul></section>")

    if report.original_lesson_material.strip():
        parts.append("<section><h2>原始教案</h2>")
        for paragraph in report.original_lesson_material.splitlines():
            text = paragraph.strip()
            if text:
                parts.append(f"<p>{escape(text)}</p>")
        parts.append("</section>")

    parts.append("<section><h2>修订后教案</h2>")
    revised = format_revised_lesson_plan(report.revised_lesson_plan)
    for paragraph in revised.splitlines():
        text = paragraph.strip()
        if text:
            parts.append(f"<p>{escape(text)}</p>")
    parts.append("</section></body></html>")
    return "".join(parts)


def to_docx_revised_plan_bytes(report: SimulationReport) -> bytes:
    from docx import Document

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc = Document()
    doc.add_heading("AI修订后教案", level=1)
    doc.add_paragraph(f"生成时间: {now}")
    doc.add_paragraph(f"学科: {report.subject}")
    doc.add_paragraph(f"课题: {report.lesson_topic}")
    doc.add_paragraph(f"年级: {report.grade}")

    if report.lesson_plan_change_summary:
        doc.add_heading("教案修改摘要", level=2)
        for item in report.lesson_plan_change_summary:
            doc.add_paragraph(item, style="List Bullet")

    if report.original_lesson_material.strip():
        doc.add_heading("原始教案", level=2)
        for paragraph in report.original_lesson_material.splitlines():
            text = paragraph.strip()
            if text:
                doc.add_paragraph(text)

    doc.add_heading("修订后教案", level=2)
    revised = format_revised_lesson_plan(report.revised_lesson_plan)
    for paragraph in revised.splitlines():
        text = paragraph.strip()
        if text:
            doc.add_paragraph(text)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_student_report_payload(report: SimulationReport, export_format: str) -> tuple[bytes | str, str, str]:
    normalized = export_format.lower()
    if normalized == "markdown":
        return to_markdown(report), "student_reaction_report.md", "text/markdown"
    if normalized == "html":
        return to_html(report), "student_reaction_report.html", "text/html"
    if normalized == "word":
        return (
            to_docx_bytes(report),
            "student_reaction_report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    raise ValueError(f"不支持的导出格式: {export_format}")


def build_revised_plan_payload(report: SimulationReport, export_format: str) -> tuple[bytes | str, str, str]:
    normalized = export_format.lower()
    if normalized == "markdown":
        return to_markdown_revised_plan(report), "revised_lesson_plan.md", "text/markdown"
    if normalized == "html":
        return to_html_revised_plan(report), "revised_lesson_plan.html", "text/html"
    if normalized == "word":
        return (
            to_docx_revised_plan_bytes(report),
            "revised_lesson_plan.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    raise ValueError(f"不支持的导出格式: {export_format}")
